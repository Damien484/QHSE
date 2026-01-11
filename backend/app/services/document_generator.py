"""
Service de génération de documents DUERP
Génère des documents PDF et DOCX conformes à la réglementation française
"""
import os
from datetime import datetime
from pathlib import Path
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


class DUERPDocumentGenerator:
    """Générateur de documents DUERP"""

    def __init__(self):
        self.output_dir = Path(__file__).resolve().parent.parent.parent.parent / 'generated_documents'
        self.output_dir.mkdir(exist_ok=True)

    def generate_pdf(self, duerp):
        """
        Génère un document PDF pour le DUERP

        Args:
            duerp: Instance du modèle DUERP

        Returns:
            str: Chemin du fichier PDF généré
        """
        # Nom du fichier
        filename = f"DUERP_{duerp.entreprise_nom.replace(' ', '_')}_{duerp.version}_{datetime.now().strftime('%Y%m%d')}.pdf"
        filepath = self.output_dir / filename

        # Création du document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            spaceBefore=12
        ))
        styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY
        ))

        # Contenu du document
        story = []

        # Page de garde
        story.extend(self._generate_cover_page(duerp, styles))
        story.append(PageBreak())

        # Sommaire et informations
        story.extend(self._generate_info_section(duerp, styles))
        story.append(PageBreak())

        # Tableau récapitulatif des risques
        story.extend(self._generate_risk_summary(duerp, styles))
        story.append(PageBreak())

        # Détail par unité de travail
        story.extend(self._generate_detailed_risks(duerp, styles))

        # Génération du PDF
        doc.build(story)

        return str(filepath)

    def _generate_cover_page(self, duerp, styles):
        """Génère la page de garde"""
        elements = []

        elements.append(Spacer(1, 3*cm))
        elements.append(Paragraph("DOCUMENT UNIQUE", styles['CustomTitle']))
        elements.append(Paragraph("D'ÉVALUATION DES RISQUES PROFESSIONNELS", styles['CustomTitle']))
        elements.append(Spacer(1, 2*cm))

        elements.append(Paragraph(f"<b>{duerp.entreprise_nom}</b>", styles['CustomTitle']))
        elements.append(Spacer(1, 1*cm))

        if duerp.entreprise_siret:
            elements.append(Paragraph(f"SIRET: {duerp.entreprise_siret}", styles['CustomNormal']))
        if duerp.entreprise_adresse:
            elements.append(Paragraph(f"Adresse: {duerp.entreprise_adresse}", styles['CustomNormal']))
        if duerp.entreprise_activite:
            elements.append(Paragraph(f"Activité: {duerp.entreprise_activite}", styles['CustomNormal']))

        elements.append(Spacer(1, 2*cm))

        # Informations du document
        info_data = [
            ['Version:', duerp.version],
            ['Date de création:', duerp.date_creation.strftime('%d/%m/%Y') if duerp.date_creation else 'N/A'],
            ['Dernière mise à jour:', duerp.date_derniere_maj.strftime('%d/%m/%Y') if duerp.date_derniere_maj else 'N/A'],
            ['Responsable évaluation:', duerp.responsable_evaluation or 'Non spécifié'],
            ['Statut:', duerp.statut.upper()]
        ]

        info_table = Table(info_data, colWidths=[6*cm, 8*cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E6E6E6'))
        ]))

        elements.append(info_table)

        return elements

    def _generate_info_section(self, duerp, styles):
        """Génère la section d'informations générales"""
        elements = []

        elements.append(Paragraph("1. INFORMATIONS GÉNÉRALES", styles['CustomHeading2']))
        elements.append(Spacer(1, 0.5*cm))

        # Contexte réglementaire
        context_text = """
        Le Document Unique d'Évaluation des Risques Professionnels (DUERP) est une obligation légale
        pour toute entreprise employant au moins un salarié (Articles L4121-1 à L4121-5 et R4121-1 à R4121-4
        du Code du travail). Il recense l'ensemble des risques pour la santé et la sécurité des travailleurs
        et les actions de prévention et de protection qui en découlent.
        """
        elements.append(Paragraph(context_text, styles['CustomNormal']))
        elements.append(Spacer(1, 0.5*cm))

        # Méthodologie
        elements.append(Paragraph("2. MÉTHODOLOGIE D'ÉVALUATION", styles['CustomHeading2']))
        elements.append(Spacer(1, 0.3*cm))

        methodology_text = """
        L'évaluation des risques a été réalisée selon la méthode de cotation suivante :<br/>
        <b>Criticité = Gravité × Probabilité</b><br/><br/>

        <b>Gravité (G):</b> 1 = Mineure, 2 = Moyenne, 3 = Grave, 4 = Très grave<br/>
        <b>Probabilité (P):</b> 1 = Très improbable, 2 = Improbable, 3 = Probable, 4 = Très probable<br/><br/>

        <b>Niveaux de risque:</b><br/>
        - Acceptable (1-2): Risque faible, surveillance normale<br/>
        - Modéré (3-6): Risque modéré, actions de prévention à planifier<br/>
        - Important (8-12): Risque important, actions de prévention prioritaires<br/>
        - Critique (16): Risque critique, actions immédiates requises
        """
        elements.append(Paragraph(methodology_text, styles['CustomNormal']))

        return elements

    def _generate_risk_summary(self, duerp, styles):
        """Génère le tableau récapitulatif des risques"""
        elements = []

        elements.append(Paragraph("3. TABLEAU RÉCAPITULATIF DES RISQUES", styles['CustomHeading2']))
        elements.append(Spacer(1, 0.5*cm))

        # Calculer les statistiques
        total_risques = 0
        risques_critiques = 0
        risques_importants = 0
        risques_moderes = 0
        risques_acceptables = 0

        for unite in duerp.unites_travail:
            for risque in unite.risques:
                total_risques += 1
                if risque.niveau_risque == 'Critique':
                    risques_critiques += 1
                elif risque.niveau_risque == 'Important':
                    risques_importants += 1
                elif risque.niveau_risque == 'Modéré':
                    risques_moderes += 1
                else:
                    risques_acceptables += 1

        # Tableau de statistiques
        stats_data = [
            ['<b>Niveau de risque</b>', '<b>Nombre</b>', '<b>%</b>'],
            ['Risques critiques', str(risques_critiques), f"{(risques_critiques/total_risques*100 if total_risques > 0 else 0):.1f}%"],
            ['Risques importants', str(risques_importants), f"{(risques_importants/total_risques*100 if total_risques > 0 else 0):.1f}%"],
            ['Risques modérés', str(risques_moderes), f"{(risques_moderes/total_risques*100 if total_risques > 0 else 0):.1f}%"],
            ['Risques acceptables', str(risques_acceptables), f"{(risques_acceptables/total_risques*100 if total_risques > 0 else 0):.1f}%"],
            ['<b>Total</b>', f'<b>{total_risques}</b>', '<b>100%</b>']
        ]

        stats_table = Table(stats_data, colWidths=[8*cm, 3*cm, 3*cm])
        stats_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -2), 'Helvetica'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#FF6B6B')),
            ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#FFA500')),
            ('BACKGROUND', (0, 3), (-1, 3), colors.HexColor('#FFD700')),
            ('BACKGROUND', (0, 4), (-1, 4), colors.HexColor('#90EE90')),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#E6E6E6'))
        ]))

        elements.append(stats_table)
        elements.append(Spacer(1, 1*cm))

        # Tableau par unité de travail
        if duerp.unites_travail:
            elements.append(Paragraph("Répartition par unité de travail:", styles['CustomNormal']))
            elements.append(Spacer(1, 0.3*cm))

            unite_data = [['<b>Unité de travail</b>', '<b>Nombre de risques</b>', '<b>Risques critiques/importants</b>']]

            for unite in duerp.unites_travail:
                nb_risques = len(unite.risques)
                nb_critiques = len([r for r in unite.risques if r.niveau_risque in ['Critique', 'Important']])
                unite_data.append([unite.nom, str(nb_risques), str(nb_critiques)])

            unite_table = Table(unite_data, colWidths=[8*cm, 4*cm, 5*cm])
            unite_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')])
            ]))

            elements.append(unite_table)

        return elements

    def _generate_detailed_risks(self, duerp, styles):
        """Génère le détail des risques par unité de travail"""
        elements = []

        elements.append(Paragraph("4. ÉVALUATION DÉTAILLÉE DES RISQUES", styles['CustomHeading2']))
        elements.append(Spacer(1, 0.5*cm))

        for unite in duerp.unites_travail:
            # Titre de l'unité
            elements.append(Paragraph(f"<b>Unité de travail: {unite.nom}</b>", styles['CustomHeading2']))

            if unite.description:
                elements.append(Paragraph(f"Description: {unite.description}", styles['CustomNormal']))
            if unite.localisation:
                elements.append(Paragraph(f"Localisation: {unite.localisation}", styles['CustomNormal']))
            if unite.nombre_employes:
                elements.append(Paragraph(f"Nombre d'employés: {unite.nombre_employes}", styles['CustomNormal']))

            elements.append(Spacer(1, 0.3*cm))

            # Tableau des risques
            if unite.risques:
                for idx, risque in enumerate(unite.risques, 1):
                    # Couleur selon le niveau de risque
                    color_map = {
                        'Critique': colors.HexColor('#FF6B6B'),
                        'Important': colors.HexColor('#FFA500'),
                        'Modéré': colors.HexColor('#FFD700'),
                        'Acceptable': colors.HexColor('#90EE90')
                    }
                    risk_color = color_map.get(risque.niveau_risque, colors.white)

                    # En-tête du risque
                    risk_header = [
                        [f'<b>Risque #{idx}</b>', f'<b>{risque.categorie}</b>', f'<b>Criticité: {risque.criticite} - {risque.niveau_risque}</b>']
                    ]

                    risk_header_table = Table(risk_header, colWidths=[3*cm, 7*cm, 7*cm])
                    risk_header_table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BACKGROUND', (0, 0), (-1, -1), risk_color),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                    ]))

                    elements.append(risk_header_table)

                    # Détails du risque
                    risk_details = [
                        ['Description:', risque.description or 'N/A'],
                        ['Situation de danger:', risque.situation_danger or 'N/A'],
                        ['Gravité:', f"{risque.gravite}/4"],
                        ['Probabilité:', f"{risque.probabilite}/4"],
                        ['Fréquence exposition:', risque.frequence_exposition or 'N/A'],
                        ['Personnes exposées:', f"{risque.personnes_exposees or 0} - {risque.personnes_concernees or 'N/A'}"]
                    ]

                    risk_table = Table(risk_details, colWidths=[5*cm, 12*cm])
                    risk_table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F5F5F5'))
                    ]))

                    elements.append(risk_table)

                    # Mesures de prévention
                    if risque.mesures_prevention:
                        mesures_data = [['<b>Type</b>', '<b>Description</b>', '<b>Statut</b>', '<b>Responsable</b>']]

                        for mesure in risque.mesures_prevention:
                            mesures_data.append([
                                mesure.type_mesure or 'N/A',
                                mesure.description or 'N/A',
                                mesure.statut or 'N/A',
                                mesure.responsable or 'N/A'
                            ])

                        mesures_table = Table(mesures_data, colWidths=[4*cm, 7*cm, 3*cm, 3*cm])
                        mesures_table.setStyle(TableStyle([
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                            ('FONTSIZE', (0, 0), (-1, -1), 8),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke)
                        ]))

                        elements.append(Paragraph("<b>Mesures de prévention:</b>", styles['CustomNormal']))
                        elements.append(mesures_table)

                    elements.append(Spacer(1, 0.5*cm))
            else:
                elements.append(Paragraph("Aucun risque identifié pour cette unité.", styles['CustomNormal']))

            elements.append(Spacer(1, 0.5*cm))
            elements.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
            elements.append(Spacer(1, 0.5*cm))

        return elements

    def generate_docx(self, duerp):
        """
        Génère un document DOCX pour le DUERP

        Args:
            duerp: Instance du modèle DUERP

        Returns:
            str: Chemin du fichier DOCX généré
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Nom du fichier
            filename = f"DUERP_{duerp.entreprise_nom.replace(' ', '_')}_{duerp.version}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = self.output_dir / filename

            # Création du document
            document = Document()

            # Page de garde
            title = document.add_heading('DOCUMENT UNIQUE', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = document.add_heading("D'ÉVALUATION DES RISQUES PROFESSIONNELS", level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()
            company = document.add_heading(duerp.entreprise_nom, level=1)
            company.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Informations entreprise
            if duerp.entreprise_siret:
                p = document.add_paragraph(f"SIRET: {duerp.entreprise_siret}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if duerp.entreprise_adresse:
                p = document.add_paragraph(f"Adresse: {duerp.entreprise_adresse}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if duerp.entreprise_activite:
                p = document.add_paragraph(f"Activité: {duerp.entreprise_activite}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_page_break()

            # Informations du document
            document.add_heading('INFORMATIONS GÉNÉRALES', level=1)

            table = document.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'

            info_data = [
                ('Version:', duerp.version),
                ('Date de création:', duerp.date_creation.strftime('%d/%m/%Y') if duerp.date_creation else 'N/A'),
                ('Dernière mise à jour:', duerp.date_derniere_maj.strftime('%d/%m/%Y') if duerp.date_derniere_maj else 'N/A'),
                ('Responsable évaluation:', duerp.responsable_evaluation or 'Non spécifié'),
                ('Statut:', duerp.statut.upper())
            ]

            for i, (label, value) in enumerate(info_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)

            document.add_page_break()

            # Risques détaillés
            document.add_heading('ÉVALUATION DÉTAILLÉE DES RISQUES', level=1)

            for unite in duerp.unites_travail:
                document.add_heading(f"Unité: {unite.nom}", level=2)

                if unite.description:
                    document.add_paragraph(f"Description: {unite.description}")
                if unite.localisation:
                    document.add_paragraph(f"Localisation: {unite.localisation}")

                for idx, risque in enumerate(unite.risques, 1):
                    document.add_heading(f"Risque #{idx}: {risque.categorie}", level=3)

                    risk_table = document.add_table(rows=7, cols=2)
                    risk_table.style = 'Light List Accent 1'

                    risk_info = [
                        ('Description:', risque.description or 'N/A'),
                        ('Situation de danger:', risque.situation_danger or 'N/A'),
                        ('Gravité:', f"{risque.gravite}/4"),
                        ('Probabilité:', f"{risque.probabilite}/4"),
                        ('Criticité:', f"{risque.criticite} - {risque.niveau_risque}"),
                        ('Fréquence exposition:', risque.frequence_exposition or 'N/A'),
                        ('Personnes exposées:', f"{risque.personnes_exposees or 0}")
                    ]

                    for i, (label, value) in enumerate(risk_info):
                        risk_table.rows[i].cells[0].text = label
                        risk_table.rows[i].cells[1].text = str(value)

                    # Mesures de prévention
                    if risque.mesures_prevention:
                        document.add_paragraph('Mesures de prévention:', style='Intense Quote')

                        mesures_table = document.add_table(rows=len(risque.mesures_prevention) + 1, cols=4)
                        mesures_table.style = 'Light Grid Accent 1'

                        # En-tête
                        headers = ['Type', 'Description', 'Statut', 'Responsable']
                        for i, header in enumerate(headers):
                            mesures_table.rows[0].cells[i].text = header

                        # Données
                        for i, mesure in enumerate(risque.mesures_prevention, 1):
                            mesures_table.rows[i].cells[0].text = mesure.type_mesure or 'N/A'
                            mesures_table.rows[i].cells[1].text = mesure.description or 'N/A'
                            mesures_table.rows[i].cells[2].text = mesure.statut or 'N/A'
                            mesures_table.rows[i].cells[3].text = mesure.responsable or 'N/A'

                    document.add_paragraph()

            # Sauvegarde
            document.save(str(filepath))

            return str(filepath)

        except ImportError:
            raise Exception("python-docx n'est pas installé. Installez-le avec: pip install python-docx")
